from langchain_openai import OpenAI  # Import from langchain-openai now
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.agents import initialize_agent, Tool
from langchain.tools import StructuredTool
from dotenv import load_dotenv
import os
from docx import Document  # Import for creating Word documents

# Load environment variables from .env file
load_dotenv()

# Now you can access the OPEN_AI_API_KEY from the environment variables
OPEN_AI_API_KEY = os.getenv('OPEN_AI_API_KEY')

# Initialize the LLM (Large Language Model) with your API key
llm = OpenAI(temperature=0.7, openai_api_key=OPEN_AI_API_KEY)

# Define a tool to gather the requirements
def gather_business_requirements(inputs):
    print(f"Inputs in gather_business_requirements: {inputs}")
    
    # Ensure inputs is a dictionary and not a string
    if isinstance(inputs, str):
        inputs = eval(inputs)  # Convert string back to dictionary if necessary

    business_goals = inputs.get("business_goals", "")
    user_stories = inputs.get("user_stories", "")
    acceptance_criteria = inputs.get("acceptance_criteria", "")
    stakeholders = inputs.get("stakeholders", "")
    
    document = f"""
    # Business Requirements
    ## Business Goals:
    {business_goals}
    
    ## User Stories:
    {user_stories}
    
    ## Acceptance Criteria:
    {acceptance_criteria}
    
    ## Stakeholders:
    {stakeholders}
    """
    return document

# Define a tool to generate User Stories
def generate_user_stories(inputs):
    print(f"Inputs in generate_user_stories: {inputs}")
    
    # Ensure inputs is a dictionary and not a string
    if isinstance(inputs, str):
        inputs = eval(inputs)  # Convert string back to dictionary if necessary

    user_stories = inputs.get("user_stories", "")
    prompt_template = "As a {user_role}, I want {functionality} so that {benefit}"
    stories = "\n".join([prompt_template.format(**story) for story in user_stories])
    return stories

def generate_functional_requirements(inputs):
    print(f"Inputs in generate_functional_requirements: {inputs}")

    # Ensure inputs is a dictionary and handle any unexpected structure
    if isinstance(inputs, str):
        try:
            inputs = eval(inputs)  # Convert string back to dictionary
            print(f"Inputs after eval: {inputs}")
        except Exception as e:
            print(f"Error while evaluating inputs: {e}")
            return "Error processing inputs"

    # Check if the expected keys exist in the inputs
    if not isinstance(inputs, dict):
        return "Error: Inputs should be a dictionary."

    # Extract required values with default fallbacks
    business_goals = inputs.get("business_goals", "No business goals provided.")
    acceptance_criteria = inputs.get("acceptance_criteria", "No acceptance criteria provided.")

    # Debugging print statements to confirm the extracted values
    print(f"Business Goals: {business_goals}")
    print(f"Acceptance Criteria: {acceptance_criteria}")

    # Create the document
    document = f"""
    ## Functional Requirements:
    {business_goals}
    
    ## Acceptance Criteria:
    {acceptance_criteria}
    """
    
    # Final debug print to confirm the document structure
    print(f"Generated Functional Requirements:\n{document}")
    
    return document

# Define a tool to create Stakeholder Analysis
def stakeholder_analysis(inputs):
    print(f"Inputs in stakeholder_analysis: {inputs}")
    
    # Ensure inputs is a dictionary and not a string
    if isinstance(inputs, str):
        inputs = eval(inputs)  # Convert string back to dictionary if necessary

    stakeholders = inputs.get("stakeholders", "")
    report = ""
    for stakeholder in stakeholders:
        report += f"### Stakeholder: {stakeholder['name']}\nRole: {stakeholder['role']}\nNeeds: {stakeholder['needs']}\nPain Points: {stakeholder['pain_points']}\n\n"
    return report

# Define a function to generate a concise summary for the Project Manager
def generate_summary_for_pm(inputs):
    print("\nGenerating summary for Project Manager...")
    business_goals = inputs.get("business_goals", "")
    user_stories = inputs.get("user_stories", "")
    acceptance_criteria = inputs.get("acceptance_criteria", "")
    stakeholders = inputs.get("stakeholders", "")
    
    summary = f"""
    ### **Project Summary for Review** 📋

    **Business Overview**
    - **Business Goals** 🌱:
        {business_goals}

    **User Stories** 📖:
    """
    for story in user_stories:
        summary += f"- **{story['user_role']}**: As a {story['user_role']}, I want {story['functionality']} so that {story['benefit']}\n"
    
    summary += f"""
    **Acceptance Criteria** ✅:
    {acceptance_criteria}
    
    **Stakeholder Analysis** 🤝:
    """
    for stakeholder in stakeholders:
        summary += f"""
        - **{stakeholder['name']}**:  
            - Role: {stakeholder['role']}
            - Needs: {stakeholder['needs']}
            - Pain Points: {stakeholder['pain_points']}
        """
    
    summary += f"""
    **Next Steps for Feedback**
    - Review the business goals and user stories to ensure they align with the project objectives.
    - Provide feedback on any adjustments needed for the functional requirements and acceptance criteria.
    - Clarify any additional requirements or challenges with stakeholders, especially around tools and resources.

    ### **Feedback Section**
    - What to improve or adjust: __[Space for feedback]__
    - Additional considerations: __[Space for feedback]__
    """
    
    print(f"\nGenerated Project Manager Summary:\n{summary}")
    return summary

# Function to create a Word document from the summary
def create_word_document(summary, filename="Project_Summary.docx"):
    # Create a new Document object
    doc = Document()
    
    # Add a title
    doc.add_heading('Project Management Summary', 0)
    
    # Add the content from the summary
    doc.add_paragraph(summary)
    
    # Save the document
    doc.save(filename)
    
    print(f"\nWord document '{filename}' has been generated successfully!")

# Create and initialize agent with tools
tools = [
    Tool(func=gather_business_requirements, name="GatherBusinessRequirements", description="Gathers business requirements, user stories, and stakeholder information."),
    Tool(func=generate_user_stories, name="GenerateUserStories", description="Generates user stories from the provided input."),
    Tool(func=generate_functional_requirements, name="GenerateFunctionalRequirements", description="Generates functional requirements document based on input."),
    Tool(func=stakeholder_analysis, name="StakeholderAnalysis", description="Creates a stakeholder analysis report based on provided stakeholder details.")
]

agent = initialize_agent(tools, llm, agent_type="zero_shot", verbose=True)

# Inputs before invoking agent
inputs_with_input_key = {
    "input": {
        "business_goals": "To expand the product line and increase customer reach.",
        "acceptance_criteria": "All product categories should be easy to navigate and marketing campaigns should result in a 10% increase in engagement.",
        "user_stories": [
            {"user_role": "product manager", "functionality": "create new product categories", "benefit": "expand product options for customers"},
            {"user_role": "marketing manager", "functionality": "launch marketing campaigns", "benefit": "increase brand visibility"}
        ],
        "stakeholders": [
            {"name": "Alice", "role": "Product Manager", "needs": "Clear roadmap for new product features", "pain_points": "Slow development cycles"},
            {"name": "Bob", "role": "Marketing Manager", "needs": "Timely data on campaign performance", "pain_points": "Ineffective tracking tools"}
        ]
    }
}

# Print inputs before invoking the agent
print(f"Inputs before invoking agent: {inputs_with_input_key}")

# Run the agent with the corrected inputs structure
output = agent.invoke(inputs_with_input_key)

# Now generate a concise summary for the Project Manager
summary = generate_summary_for_pm(inputs_with_input_key['input'])

# Print the generated summary for the Project Manager
print("\nFinal Summary for Project Manager: ")
print(summary)

print(output)

# Create the Word document with the Project Manager summary
create_word_document(summary)