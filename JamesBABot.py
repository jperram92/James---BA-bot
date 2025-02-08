from langchain_openai import OpenAI  # Import from langchain-openai now
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.agents import initialize_agent, Tool
from langchain.tools import StructuredTool
from dotenv import load_dotenv
import os
from docx import Document  # Import for creating Word documents
from docx.shared import RGBColor

# Load environment variables from .env file
load_dotenv()

# Now you can access the OPEN_AI_API_KEY from the environment variables
OPEN_AI_API_KEY = os.getenv('OPEN_AI_API_KEY')

# Initialize the LLM (Large Language Model) with your API key
llm = OpenAI(temperature=0.7, openai_api_key=OPEN_AI_API_KEY)

# Define a tool to gather the requirements
def gather_business_requirements(inputs):
    print(f"Inputs in gather_business_requirements: {inputs}")
    
    # Ensure inputs is a dictionary and not a string
    if isinstance(inputs, str):
        inputs = eval(inputs)  # Convert string back to dictionary if necessary

    business_goals = inputs.get("business_goals", "")
    user_stories = inputs.get("user_stories", "")
    acceptance_criteria = inputs.get("acceptance_criteria", "")
    stakeholders = inputs.get("stakeholders", "")
    
    document = f"""
    # Business Requirements
    ## Business Goals:
    {business_goals}
    
    ## User Stories:
    {user_stories}
    
    ## Acceptance Criteria:
    {acceptance_criteria}
    
    ## Stakeholders:
    {stakeholders}
    """
    return document

# Define a tool to generate User Stories
def generate_user_stories(inputs):
    print(f"Inputs in generate_user_stories: {inputs}")
    
    # Ensure inputs is a dictionary and not a string
    if isinstance(inputs, str):
        inputs = eval(inputs)  # Convert string back to dictionary if necessary

    user_stories = inputs.get("user_stories", "")
    prompt_template = "As a {user_role}, I want {functionality} so that {benefit}"
    stories = "\n".join([prompt_template.format(**story) for story in user_stories])
    return stories

def generate_functional_requirements(inputs):
    print(f"Inputs in generate_functional_requirements: {inputs}")

    # Ensure inputs is a dictionary and handle any unexpected structure
    if isinstance(inputs, str):
        try:
            inputs = eval(inputs)  # Convert string back to dictionary
            print(f"Inputs after eval: {inputs}")
        except Exception as e:
            print(f"Error while evaluating inputs: {e}")
            return "Error processing inputs"

    # Check if the expected keys exist in the inputs
    if not isinstance(inputs, dict):
        return "Error: Inputs should be a dictionary."

    # Extract required values with default fallbacks
    business_goals = inputs.get("business_goals", "No business goals provided.")
    acceptance_criteria = inputs.get("acceptance_criteria", "No acceptance criteria provided.")

    # Debugging print statements to confirm the extracted values
    print(f"Business Goals: {business_goals}")
    print(f"Acceptance Criteria: {acceptance_criteria}")

    # Create the document
    document = f"""
    ## Functional Requirements:
    {business_goals}
    
    ## Acceptance Criteria:
    {acceptance_criteria}
    """
    
    # Final debug print to confirm the document structure
    print(f"Generated Functional Requirements:\n{document}")
    
    return document

# Define a tool to create Stakeholder Analysis
def stakeholder_analysis(inputs):
    print(f"Inputs in stakeholder_analysis: {inputs}")
    
    # Ensure inputs is a dictionary and not a string
    if isinstance(inputs, str):
        inputs = eval(inputs)  # Convert string back to dictionary if necessary

    stakeholders = inputs.get("stakeholders", "")
    report = ""
    for stakeholder in stakeholders:
        report += f"### Stakeholder: {stakeholder['name']}\nRole: {stakeholder['role']}\nNeeds: {stakeholder['needs']}\nPain Points: {stakeholder['pain_points']}\n\n"
    return report

# Define a function to generate a concise summary for the Project Manager
def generate_summary_for_pm(inputs):
    print("\nGenerating summary for Project Manager...")
    business_goals = inputs.get("business_goals", "")
    user_stories = inputs.get("user_stories", "")
    acceptance_criteria = inputs.get("acceptance_criteria", "")
    stakeholders = inputs.get("stakeholders", "")
    
    summary = f"""
    ### **Project Summary for Review** 📋

    **Business Overview**
    - **Business Goals** 🌱:
        {business_goals}

    **User Stories** 📖:
    """
    for story in user_stories:
        summary += f"- **{story['user_role']}**: As a {story['user_role']}, I want {story['functionality']} so that {story['benefit']}\n"
    
    summary += f"""
    **Acceptance Criteria** ✅:
    {acceptance_criteria}
    
    **Stakeholder Analysis** 🤝:
    """
    for stakeholder in stakeholders:
        summary += f"""
        - **{stakeholder['name']}**:  
            - Role: {stakeholder['role']}
            - Needs: {stakeholder['needs']}
            - Pain Points: {stakeholder['pain_points']}
        """
    
    summary += f"""
    **Next Steps for Feedback**
    - Review the business goals and user stories to ensure they align with the project objectives.
    - Provide feedback on any adjustments needed for the functional requirements and acceptance criteria.
    - Clarify any additional requirements or challenges with stakeholders, especially around tools and resources.

    ### **Feedback Section**
    - What to improve or adjust: __[Space for feedback]__
    - Additional considerations: __[Space for feedback]__
    """
    
    print(f"\nGenerated Project Manager Summary:\n{summary}")
    return summary

# Function to create a Word document from the summary
def create_word_document(summary, filename="Project_Summary.docx"):
    # Create a new Document object
    doc = Document()
    
    # Add a title
    doc.add_heading('Project Management Summary', 0)
    
    # Add the content from the summary
    doc.add_paragraph(summary)
    
    # Save the document
    doc.save(filename)
    
    print(f"\nWord document '{filename}' has been generated successfully!")

from docx import Document
from docx.shared import RGBColor

# Function to generate the project management summary document in the format similar to the uploaded image
def create_project_management_summary(summary, filename="Project_Management_Summary.docx"):
    # Create a new Document object
    doc = Document()
    
    # Add a title with a logo-like green style
    doc.add_heading('Techno-PM', 0).alignment = 1  # Center-aligned for a professional look
    run = doc.paragraphs[0].runs[0]
    run.font.size = 240000  # Adjust the font size to make it prominent
    
    # Add "Project Management Templates" in a smaller font and green
    doc.add_paragraph('Project Management Templates', style='Normal').runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    # Add Key Highlights Section
    doc.add_heading('Key Highlights', level=1)
    doc.add_paragraph('Overall the project is 25% complete.')
    doc.add_paragraph('Requirements have been delayed due to customer changes.').runs[0].bold = True
    doc.add_paragraph('Project Build is at 75%.')
    doc.add_paragraph("John on leave next week and Friday is a public holiday.")
    
    # Add Task/Issue Description Table
    doc.add_heading('Task / Issue Description', level=1)

    # Create table headers and fill in data based on user story inputs
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Task / Issue Description'
    hdr_cells[2].text = 'Owner'
    hdr_cells[3].text = 'Start'
    hdr_cells[4].text = 'Due'
    hdr_cells[5].text = 'Status'
    
    # Adding example data based on prior information; tasks, owners, start and due dates
    tasks_data = [
        [1, "Complete the Business Requirements and handover to John", "Julie", "12-Jan", "19-Mar", "WIP"],
        [2, "Organize meeting with customer to understand changes to our policies", "Sam", "13-Jan", "14-Jan", "WIP"],
        [3, "Organize system testing for the change required", "John", "11-Jan", "12-Jan", "Late"],
        [4, "Discuss the possible changes to the scope", "Kylie", "11-Feb", "14-Feb", "Close"]
    ]
    
    # Filling the table with tasks
    for task in tasks_data:
        row_cells = table.add_row().cells
        for i, task_item in enumerate(task):
            row_cells[i].text = str(task_item)
    
    # Add Comments/Updates section for each task
    doc.add_heading('Comments / Updates', level=2)
    doc.add_paragraph("15-Jan by John: The task is not complete as the business requirements are not clear.")
    doc.add_paragraph("18-Feb by John: The task is back on track after requirements have been received.")
    doc.add_paragraph("11-Jan by Sam: This meeting will need to be moved as the changes are not finalised.")
    doc.add_paragraph("13-Jan by John: Task Delayed as no capacity present.")
    doc.add_paragraph("12-Jan by Kylie: Task complete.")
    
    # Save the document
    doc.save(filename)
    return filename

# Generate the project management summary document
filename = create_project_management_summary(summary="Project Summary", filename="Project_Management_Summary.docx")

# Create and initialize agent with tools
tools = [
    Tool(func=gather_business_requirements, name="GatherBusinessRequirements", description="Gathers business requirements, user stories, and stakeholder information."),
    Tool(func=generate_user_stories, name="GenerateUserStories", description="Generates user stories from the provided input."),
    Tool(func=generate_functional_requirements, name="GenerateFunctionalRequirements", description="Generates functional requirements document based on input."),
    Tool(func=stakeholder_analysis, name="StakeholderAnalysis", description="Creates a stakeholder analysis report based on provided stakeholder details.")
]

agent = initialize_agent(tools, llm, agent_type="zero_shot", verbose=True)

# Inputs before invoking agent
inputs_with_input_key = {
    "input": {
        "business_goals": "To expand the product line and increase customer reach.",
        "acceptance_criteria": "All product categories should be easy to navigate and marketing campaigns should result in a 10% increase in engagement.",
        "user_stories": [
            {"user_role": "product manager", "functionality": "create new product categories", "benefit": "expand product options for customers"},
            {"user_role": "marketing manager", "functionality": "launch marketing campaigns", "benefit": "increase brand visibility"}
        ],
        "stakeholders": [
            {"name": "Alice", "role": "Product Manager", "needs": "Clear roadmap for new product features", "pain_points": "Slow development cycles"},
            {"name": "Bob", "role": "Marketing Manager", "needs": "Timely data on campaign performance", "pain_points": "Ineffective tracking tools"}
        ]
    }
}

# Print inputs before invoking the agent
print(f"Inputs before invoking agent: {inputs_with_input_key}")

# Run the agent with the corrected inputs structure
output = agent.invoke(inputs_with_input_key)

# Now generate a concise summary for the Project Manager
summary = generate_summary_for_pm(inputs_with_input_key['input'])

# Print the generated summary for the Project Manager
print("\nFinal Summary for Project Manager: ")
print(summary)

print(output)

# Create the Word document with the Project Manager summary
create_word_document(summary)